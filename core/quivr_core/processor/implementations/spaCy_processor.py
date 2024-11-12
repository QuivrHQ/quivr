import logging
import os
import spacy
import aiofiles
import pandas as pd
import fitz  # PyMuPDF for PDF processing
import docx  # python-docx for DOCX processing
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter, TextSplitter

from quivr_core.files.file import QuivrFile
from quivr_core.processor.processor_base import ProcessorBase
from quivr_core.processor.registry import FileExtension
from quivr_core.processor.splitter import SplitterConfig

logger = logging.getLogger("quivr_core")


class SpaCyProcessor(ProcessorBase):
    """
    SpaCyProcessor for handling various text file types with spaCy NLP.

    It extracts and processes text content using spaCy's NLP pipeline.

    ## Installation
    ```bash
    pip install spacy pandas pymupdf python-docx
    python -m spacy download en_core_web_sm
    ```
    """

    supported_extensions = [
        FileExtension.pdf,
        FileExtension.docx,
        FileExtension.txt,
        FileExtension.csv,
    ]

    def __init__(
        self,
        splitter: TextSplitter | None = None,
        splitter_config: SplitterConfig = SplitterConfig(),
        spacy_model: str = "en_core_web_sm"
    ) -> None:
        # Load spaCy model
        try:
            self.nlp = spacy.load(spacy_model)
        except Exception as e:
            logger.error(f"Failed to load spaCy model '{spacy_model}': {e}")
            raise

        self.splitter_config = splitter_config

        if splitter:
            self.text_splitter = splitter
        else:
            self.text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                chunk_size=splitter_config.chunk_size,
                chunk_overlap=splitter_config.chunk_overlap,
            )

    @property
    def processor_metadata(self):
        return {
            "processor_cls": "SpaCyProcessor",
            "chunk_overlap": self.splitter_config.chunk_overlap,
        }

    async def process_file_inner(self, file: QuivrFile) -> list[Document]:
        # Extract text based on file type
        try:
            if file.extension == FileExtension.pdf:
                text = await self.extract_text_from_pdf(file)
            elif file.extension == FileExtension.docx:
                text = await self.extract_text_from_docx(file)
            elif file.extension == FileExtension.txt:
                text = await self.extract_text_from_txt(file)
            elif file.extension == FileExtension.csv:
                text = await self.extract_text_from_csv(file)
            else:
                raise ValueError(f"Unsupported file type: {file.extension}")
            
            # Check for empty text
            if not text:
                logger.warning(f"No content extracted from file: {file.path}")
                return []

            # Apply spaCy NLP processing
            doc = Document(page_content=text)
            processed_docs = self.text_splitter.split_documents([doc])

            for doc in processed_docs:
                spacy_doc = self.nlp(doc.page_content)
                doc.metadata.update({
                    "chunk_size": len(spacy_doc),
                    "entities": [(ent.text, ent.label_) for ent in spacy_doc.ents],
                    "sentences": [sent.text for sent in spacy_doc.sents]
                })
                doc.page_content = spacy_doc.text

            return processed_docs

        except Exception as e:
            logger.error(f"Error processing file '{file.path}': {e}")
            return []

    async def extract_text_from_pdf(self, file: QuivrFile) -> str:
        try:
            async with file.open():
                doc = fitz.open(file.path)
                text = ""
                for page in doc:
                    text += page.get_text()
                return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF '{file.path}': {e}")
            return ""

    async def extract_text_from_docx(self, file: QuivrFile) -> str:
        try:
            doc = docx.Document(file.path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX '{file.path}': {e}")
            return ""

    async def extract_text_from_txt(self, file: QuivrFile) -> str:
        try:
            async with aiofiles.open(file.path, mode="r") as f:
                content = await f.read()
            return content
        except Exception as e:
            logger.error(f"Error extracting text from TXT '{file.path}': {e}")
            return ""

    async def extract_text_from_csv(self, file: QuivrFile) -> str:
        try:
            df = pd.read_csv(file.path)
            return ' '.join(df.astype(str).values.flatten())
        except Exception as e:
            logger.error(f"Error extracting text from CSV '{file.path}': {e}")
            return ""
